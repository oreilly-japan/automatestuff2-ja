# 演習プロジェクト 15.7.2  招待状の生成
#  guests.txtを読み込んで、Word文書を作成

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

# guests.txtを開く
guest_file = open('guests.txt', 'r', encoding='utf-8')

# 新規Word文書
doc = docx.Document()

# 1行ずつ処理する
for guest in guest_file:
    # 前後の空白や改行を削除する
    guest = guest.strip()
    if guest == '':
        continue
    doc.add_paragraph('It would be a pleasure to have the company of')
    doc.add_paragraph(guest)
    doc.add_paragraph('at 11010 Memory Lane on the Evening of')
    doc.add_paragraph('April 1st')
    doc.add_paragraph("at 7 o'clock")
    # 改ページ
    doc.add_page_break()

# すべての行をセンタリングする
for p in doc.paragraphs:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ファイルに保存
doc.save('invitation.docx')

