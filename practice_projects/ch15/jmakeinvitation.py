# 演習プロジェクト 15.7.2  招待状の生成（日本語版）
# jguests.txtを読み込んで、Word文書を作成
# 「様」などの敬称は、jguests.txtに書いておく。

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

guest_file = open('jguests.txt', 'r', encoding='utf-8')

doc = docx.Document()

for guest in guest_file:
    guest = guest.strip()
    if guest == '':
        continue
    doc.add_paragraph(guest)
    doc.add_paragraph('拝啓　時下ますますご盛栄のこととお喜び申し上げます。')
    doc.add_paragraph('このたび下記の通りパーティーを開催しますので、' +
                      'ご出席賜りますようよろしくお願い申し上げます。')
    doc.add_paragraph('敬具').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph('記').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\t日時：4月1日　19:00')
    doc.add_paragraph('\t会場：ホテル・オライリー')
    doc.add_paragraph('以上').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()

doc.save('jinvitation.docx')

